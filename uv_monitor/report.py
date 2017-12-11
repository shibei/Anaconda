from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
import public_class as public
import os


class Report(object):
    """docstring for Report"""

    def __init__(self, data, day=1, cycle=1, path=r'/root/ipython/UV分析结果/Report/'):
        # super(Report, self).__init__()
        self.data = data
        self.day = day
        self.cycle = cycle
        self.path =path

    def _addTab(self, df, document):
        df = df.fillna('nan')
        rowsLen = len(df)
        columnsLen = len(df.columns)
        columns = df.columns
        table = document.add_table(rows=1, cols=columnsLen+1)
        hdr_cells = table.rows[0].cells
        for t in range(columnsLen):
            hdr_cells[t+1].text = str(columns[t])
        for r in range(rowsLen):
            row = df.loc[df.index[r]].values
            row_cells = table.add_row().cells
            row_cells[0].text = str(df.index[r])
            for t in range(columnsLen):
                if str(df.index[r]) == 'troubleRate' or t == 7:
                    row_cells[t+1].text = str(row[t]*100)[:4]+r'%'
                else:
                    row_cells[t+1].text = str(row[t])

    def _reFormatData(self):
        info = self.data['troubleTotal']
        uvDay1 = self.data['troubleUvDay1']
        uvDay2 = self.data['troubleUvDay2']
        rate = self.data['troubleRate']
        out = {}
        for id in info.index:
            out[id] = {}
            out[id]['Total'] = pd.DataFrame(info.drop_duplicates().loc[id])
            hosUvData = np.transpose(pd.DataFrame(uvDay1.loc[id][1:]))
            hosUvData = hosUvData.append(uvDay2.loc[id][1:])
            hosUvData = hosUvData.append(rate.loc[id][1:])
            hosUvData.insert(0, column='title', value=[
                             'troubleUvDay1', 'troubleUvDay2', 'troubleRate'])
            hosUvData = hosUvData.set_index('title')
            out[id]['Data'] = hosUvData
        return out

    def makeReport(self):
        day = self.day
        cycle = self.cycle
        path = self.path
        out = self.data
        data = self._reFormatData()
        mainTitle = r'%s周对比分析——%s' % (public.chineseNumber(
            cycle), str(public.getDate(day)).replace('-', ''))
        paragraph = r'以下为%s的%s周分析报告。' % (
            str(public.getDate(day)).replace('-', ''), public.chineseNumber(cycle))
        total = out['troubleTotal'][
            [r'故障类型', r'医院名称', r'医院级别', r'报修问题']].drop_duplicates(r'医院名称')
        floder = path + r'%s周对比/' % public.chineseNumber(cycle)
        if not os.path.exists(floder):
            os.makedirs(floder)

        fileName = mainTitle
        fullPath = floder + fileName+'.docx'

        document = Document()
        # 添加主标题
        document.add_heading(mainTitle, 0)
        # 添加主标题下的内容
        p = document.add_paragraph(paragraph)
        document.add_heading(r'总体概况', level=1)
        self._addTab(df=total, document=document)
        for hos in list(data.keys()):
            hosname = data[hos]['Total'].loc['医院名称'].values[0]
            document.add_heading(hosname, level=1)
            document.add_heading('医院概况', level=2)
            self._addTab(data[hos]['Total'], document)
            document.add_heading('医院UV数据', level=2)
            # print('hosData',data[hos]['Data'])
            self._addTab(data[hos]['Data'], document)
        document.save(fullPath)
        return fullPath
